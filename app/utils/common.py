import csv
import io
import re
import base64
from typing import List, Tuple, Optional
from PyPDF2 import PdfReader
from docx import Document
from datetime import datetime

from fastapi import HTTPException, UploadFile
import httpx
from app.api.database.database_interaction import DATA_URL
from app.schemas.message import FileData, Message
from fastapi.responses import JSONResponse

def create_signature_with_doc(base_cls, docstring: str):
    """
    Dynamically create a subclass of the given base class with a custom docstring.

    Useful for modifying or documenting classes at runtime (e.g., in API schemas or DSLs).

    Args:
        base_cls (type): The base class to extend.
        docstring (str): The new docstring to apply to the dynamically created class.

    Returns:
        type: A new subclass of `base_cls` with the specified docstring.
    """
    return type(base_cls.__name__, (base_cls,), {"__doc__": docstring})

def build_error_response(code: str, message: str, status: int) -> JSONResponse:
    """
    Construct a standardized JSON error response for API endpoints.

    This helper function formats error responses according to a consistent schema
    containing an error code, message, HTTP status, and a UTC timestamp.

    Args:
        code (str): A short error code identifier (e.g., "RESOURCE_NOT_FOUND").
        message (str): A human-readable error message.
        status (int): HTTP status code (e.g., 400, 404, 500).

    Returns:
        JSONResponse: A FastAPI JSONResponse object containing the formatted error.
    
    Example:
        return build_error_response("INVALID_ID", "The provided ID is not valid.", 400)
    """
    return JSONResponse(
        status_code=status,
        content={
            "error": {
                "code": code,
                "message": message,
                "status": status,
                "timestamp": datetime.utcnow().isoformat() + "Z"
            }
        }
    )


# Helper function to convert MongoDB document (_id) into a serializable dictionary
def serialize_mongo_document(doc):
    """
    Convert a MongoDB document into a serializable dictionary for API responses.

    Replaces the `_id` field with a string `id` for frontend compatibility.

    Args:
        doc (dict): A MongoDB document.

    Returns:
        dict | None: A sanitized and serializable dictionary, or None if the input is falsy.
    """
    if not doc:
        return None
    
    doc = doc.copy()
    if "_id" in doc:
        doc["id"] = str(doc["_id"])  # Replace MongoDB's _id with stringified id
        del doc["_id"]
    return doc

def serialize_image(image) -> str:
    """Convert various image types to base64 string for JSON serialization."""

    if image is None:
        return None

    if isinstance(image, str):
        if image.startswith("data:"):
            match = re.match(r"data:.*?;base64,(.*)", image)
            return match.group(1) if match else None
        return image

    if isinstance(image, bytes):
        return base64.b64encode(image).decode('utf-8')

    if hasattr(image, 'save'):  # PIL.Image.Image
        buffer = io.BytesIO()
        image.save(buffer, format='PNG')
        return base64.b64encode(buffer.getvalue()).decode('utf-8')

    # Extract from DSPy image object with a url field
    if hasattr(image, 'url') and isinstance(image.url, str):
        if image.url.startswith("data:"):
            match = re.match(r"data:.*?;base64,(.*)", image.url)
            return match.group(1) if match else None
        return image.url

    return None

def extract_text(file_data: FileData, content_bytes: bytes) -> Optional[str]:
    """    Extract text content from various file types.    
    Args:
        file_data (FileData): Metadata about the file including its type.
        content_bytes (bytes): The raw bytes of the file content.
    Returns:
        Optional[str]: Extracted text content, or None if extraction fails.
    """
    try:
        # Plain text and markdown
        if file_data.type in {"text/plain", "text/markdown"}:
            return content_bytes.decode("utf-8", errors="ignore")

        # CSV
        elif file_data.type == "text/csv":
            try:
                decoded = content_bytes.decode("utf-8", errors="ignore")
                reader = csv.reader(io.StringIO(decoded))
                rows = [" | ".join(row) for row in reader]
                return "\n".join(rows) if rows else None
            except Exception as e:
                print(f"Failed to parse CSV {file_data.name}: {e}")
                return None

        # PDF
        elif file_data.type == "application/pdf":
            reader = PdfReader(io.BytesIO(content_bytes))
            if reader.is_encrypted:
                raise ValueError(f"PDF file '{file_data.name}' is encrypted and cannot be processed.")
            extracted = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        extracted.append(text)
                except Exception as e:
                    print(f"Failed to extract text from page {page_num} in {file_data.name}: {e}")
            return "\n".join(extracted) if extracted else None

        # DOCX
        elif file_data.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(content_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables.append(" | ".join(cells))
            all_text = paragraphs + tables
            return "\n".join(all_text) if all_text else None

    except Exception as e:
        print(f"Failed to extract text from {file_data.name}: {e}")
        return None


def classify_file(files: List[FileData]) -> Tuple[List[FileData], List[FileData]]:
    """
    Classify files into images and documents based on their MIME types.

    Args:
        files (List[FileData]): List of FileData objects to classify.

    Returns:
        Tuple[List[FileData], List[FileData]]: Two lists - one for image files, one for document files.
    """
    image_files = [f for f in files if f.type.startswith("image/")]
    doc_files = [f for f in files if f.type in ("text/plain", "application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")]
    return image_files, doc_files

async def handle_file_processing(conversation_id: str, files: List[UploadFile]) -> Message:
    """    Process uploaded files, extract text content, and return a structured message.
    Args:
        conversation_id (str): The ID of the conversation to associate with the files.
        files (List[UploadFile]): List of files uploaded by the user.
    Returns:
        Message: A structured message containing the extracted text and file metadata.
    """
    # Read file contents and metadata
    file_bytes_list = [await file.read() for file in files]
    filenames = [file.filename for file in files]
    mime_types = [file.content_type for file in files]

    # Upload files
    async with httpx.AsyncClient() as client:
        response = await client.post(
            url=f"{DATA_URL}/api/conversations/upload_file",
            params={"convo_id": conversation_id},
            files=[("files", (file.filename, content)) for file, content in zip(files, file_bytes_list)]
        )
        response.raise_for_status()
        gcs_urls = response.json().get("gcs_urls", [])

    # Prepare FileData list
    processed_files = [
        FileData(name=name, type=mtype, url=url)
        for name, mtype, url in zip(filenames, mime_types, gcs_urls)
    ]

    # Classify files
    _, document_files = classify_file(processed_files)

    # Extract text content
    extracted_texts = [
        extract_text(file, content)
        for file, content in zip(processed_files, file_bytes_list)
        if file in document_files
    ]
    extracted_texts = [text for text in extracted_texts if text]

    # Combine content
    combined_content = "\n\n".join(extracted_texts) if extracted_texts else None

    return Message(
        content=combined_content,
        files=processed_files,
        timestamp=datetime.utcnow()
    )
